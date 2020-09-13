import xlrd
import numpy as np
from 题库信息 import *
import random

config = xlrd.open_workbook("模拟卷配置.xls").sheet_by_index(0)

#首先阅读配置文件
total_excercise_num = int(config.cell_value(0,1))

#读取配置文件各知识点题目数量比例
domains_excercise_ratio = []
total_weight = 0
domains = []
for idx in range(2,20):
	try:
		domain = config.cell_value(idx,0)
		domain_weight = int(config.cell_value(idx,1))
		if len(domain) <= 1:
			break
		domains_excercise_ratio.append([domain,domain_weight])
		total_weight += domain_weight
		domains.append(domain)
	except IndexError:
		break

assign_ratio = []
for i in range(len(domains_excercise_ratio)):
	domains_excercise_ratio[i][1] = domains_excercise_ratio[i][1] / total_weight
	assign_ratio.append(domains_excercise_ratio[i][1])

#分配各知识点题目数量
assigns = np.random.multinomial(total_excercise_num,assign_ratio, size = 1)[0]
domains_excercise_num = [[domains[idx],assigns[idx]] for idx in range(len(domains))]

excercise_seleceted = []
#挑选试题
bank = Bank()
for domain,num in domains_excercise_num:
	domain_index = knowledge_domains.index(domain)
	domain_data  = bank.excecises_Bank[domain_index]
	domain_questions = domain_data.questions
	sample_questions = random.sample(domain_questions,num)
	excercise_seleceted += sample_questions

#生成模拟试卷

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm
import time
doc = Document()
title = doc.add_heading("随机模拟卷",0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for idx,question in enumerate(excercise_seleceted):
	#问题的标题
	head = doc.add_heading('')
	run  = head.add_run("Q%d" % (idx + 1))
	run.font.size = Pt(20)
	run.font.name = u'微软雅黑'
	run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

	run  = head.add_run('  知识点: %s 来源: %s' % (question.domain,question.source))
	run.font.size = Pt(11)
	run.font.name = u'宋体'
	run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

	inline_shape = doc.add_picture(question.url)
	scaler = inline_shape.height / inline_shape.width
	width = 14
	inline_shape.height = Cm(width * scaler) # 设置图片高度为4cm
	inline_shape.width = Cm(width) # 设置图片宽度为4cm
doc.save("模拟卷_%d.docx"  % (int(time.time())))
